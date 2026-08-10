#!/usr/bin/env python3
"""Genera el acuerdo de colaboración en formato Word."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

TINTA = RGBColor(0x1A, 0x1C, 0x20)
SUAVE = RGBColor(0x55, 0x5A, 0x60)
ACENTO = RGBColor(0x3E, 0x5A, 0x6E)

doc = Document()

# ---------- página y estilo base ----------
s = doc.sections[0]
s.page_width, s.page_height = Cm(21.59), Cm(27.94)  # carta
s.top_margin = s.bottom_margin = Cm(2.2)
s.left_margin = s.right_margin = Cm(2.6)

normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)
normal.font.color.rgb = TINTA
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
pf = normal.paragraph_format
pf.space_after = Pt(7)
pf.line_spacing = 1.15


def linea(p, arriba=False):
    """Regla horizontal sobre o bajo un párrafo."""
    pPr = p._p.get_or_add_pPr()
    bdr = OxmlElement("w:pBdr")
    b = OxmlElement("w:top" if arriba else "w:bottom")
    b.set(qn("w:val"), "single")
    b.set(qn("w:sz"), "6")
    b.set(qn("w:space"), "6" if arriba else "4")
    b.set(qn("w:color"), "C8CCD0")
    bdr.append(b)
    pPr.append(bdr)


def titulo(texto):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(7)
    r = p.add_run(texto.upper())
    r.font.size = Pt(9.5)
    r.font.bold = True
    r.font.color.rgb = ACENTO
    r.font.name = "Calibri"
    linea(p)
    return p


def parrafo(texto, sangria=0, espacio=7):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(sangria)
    p.paragraph_format.space_after = Pt(espacio)
    p.add_run(texto)
    return p


def vineta(texto, negrita_hasta=None):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.7)
    p.paragraph_format.space_after = Pt(4)
    r0 = p.add_run("·  ")
    r0.font.color.rgb = SUAVE
    if negrita_hasta:
        r = p.add_run(negrita_hasta)
        r.font.bold = True
        p.add_run(texto)
    else:
        p.add_run(texto)
    return p


def sub(texto):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(9)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(texto)
    r.font.bold = True
    r.font.size = Pt(10.5)
    return p


# ══════════════════════════ portada ══════════════════════════
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(2)
r = p.add_run("Acuerdo de colaboración")
r.font.size = Pt(21)
r.font.name = "Cambria"
r.font.color.rgb = TINTA

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(18)
r = p.add_run("Plataforma de evaluación psicológica")
r.font.size = Pt(12.5)
r.font.name = "Cambria"
r.font.color.rgb = SUAVE

t = doc.add_table(rows=3, cols=2)
t.alignment = WD_TABLE_ALIGNMENT.LEFT
t.columns[0].width = Cm(4.2)
t.columns[1].width = Cm(11.6)
datos = [
    ("El desarrollador", "[nombre]  ·  [correo]  ·  [teléfono]"),
    ("El cliente", "[nombre]  ·  [correo]  ·  [teléfono]"),
    ("Lugar y fecha", "[ciudad], [__ de _______ de 2026]"),
]
for fila, (k, v) in zip(t.rows, datos):
    fila.cells[0].width = Cm(4.2)
    fila.cells[1].width = Cm(11.6)
    pk = fila.cells[0].paragraphs[0]
    rk = pk.add_run(k)
    rk.font.size = Pt(9)
    rk.font.color.rgb = SUAVE
    pk.paragraph_format.space_after = Pt(3)
    pv = fila.cells[1].paragraphs[0]
    pv.add_run(v).font.size = Pt(10.5)
    pv.paragraph_format.space_after = Pt(3)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(14)
p.add_run(
    "Ambas partes acuerdan colaborar en el desarrollo de una plataforma para "
    "aplicar, calificar e interpretar pruebas psicológicas en el consultorio "
    "del cliente."
)

# ══════════════════════════ qué recibe ══════════════════════════
titulo("Qué recibe el cliente")

parrafo(
    "Al cierre del proyecto, el consultorio cuenta con un sistema propio, "
    "instalado y funcionando, con lo siguiente:"
)

sub("Las tres pruebas listas para aplicar")
vineta(
    "con sus 90 pares de frases y la hoja de perfil de los 20 factores, "
    "dibujada como la lámina impresa.",
    "PAPI: ",
)
vineta(
    "con sus dos ordenamientos de 18 enunciados y los tres axiogramas, "
    "con los niveles de desarrollo de cada indicador.",
    "Hartman: ",
)
vineta(
    "con sus cuatro bloques —del candidato y del puesto— y las gráficas "
    "que comparan ambos perfiles.",
    "MABE: ",
)
parrafo(
    "Cada prueba trae sus instrucciones tomadas del manual, disponibles "
    "durante toda la aplicación.",
    sangria=0.7,
)

sub("Calificación automática")
parrafo(
    "Se capturan las respuestas y el sistema entrega los puntajes. Sin "
    "sumas, sin tablas, sin buscar valores a mano. Antes de calcular revisa "
    "que el protocolo sea válido: si algo no cuadra —una suma que no da, "
    "seis o más disimilitudes en el Hartman— lo detiene y explica por qué.",
)

sub("Informes")
parrafo(
    "Las gráficas de perfil y un texto borrador armado con el contenido de "
    "los propios manuales. El psicólogo lo edita, agrega sus notas y lo "
    "firma. Sale en PDF, con la fecha, el responsable y la versión de "
    "claves con que se calculó. Una vez firmado ya no cambia.",
)

sub("Un panel para verlo todo")
parrafo(
    "El listado de aplicaciones con su estado, el detalle de cada resultado, "
    "el historial completo de cada participante y tres perfiles de acceso: "
    "quien captura, quien interpreta y firma, y quien administra.",
)

sub("Privacidad resuelta")
parrafo(
    "Antes de empezar cualquier prueba, el evaluado acepta el aviso de "
    "privacidad y los términos, y el sistema guarda esa aceptación con su "
    "fecha. El desarrollador redacta esos textos y los deja integrados.",
)

sub("Y para arrancar")
parrafo(
    "Instalación, respaldos automáticos, capacitación al equipo y el manual "
    "de uso.",
)

parrafo(
    "El desglose completo está en el Anexo A. Lo que aparece ahí entra en el "
    "precio; lo que no, se cotiza aparte y se acuerda antes de hacerlo.",
    espacio=4,
)

# ══════════════════════════ precio ══════════════════════════
titulo("Precio y pagos")

parrafo("$15,000 pesos más IVA, en tres pagos de $5,000 más IVA:")

t = doc.add_table(rows=3, cols=2)
t.columns[0].width = Cm(9.5)
t.columns[1].width = Cm(6.3)
pagos = [
    ("Al firmar", "$5,000 + IVA"),
    ("Al cerrar la primera etapa", "$5,000 + IVA"),
    ("Al cerrar la segunda etapa", "$5,000 + IVA"),
]
for fila, (k, v) in zip(t.rows, pagos):
    fila.cells[0].width = Cm(9.5)
    fila.cells[1].width = Cm(6.3)
    pk = fila.cells[0].paragraphs[0]
    pk.paragraph_format.space_after = Pt(3)
    pk.add_run(k)
    pv = fila.cells[1].paragraphs[0]
    pv.paragraph_format.space_after = Pt(3)
    pv.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pv.add_run(v).font.bold = True

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(9)
p.add_run(
    "Cada etapa se entrega para revisión. El cliente tiene cinco días hábiles "
    "para comentarla; sin observaciones en ese plazo se da por aceptada y el "
    "pago corre dentro de los cinco días siguientes. Un pago con más de "
    "quince días de retraso permite pausar el trabajo, avisando antes."
)

# ══════════════════════════ tiempos ══════════════════════════
titulo("Tiempos")
parrafo(
    "Tres meses, a partir de que el cliente entregue los protocolos "
    "calificados a mano y confirme las licencias de las pruebas. Si el avance "
    "se detiene por un insumo pendiente o por agendar una sesión, ese tiempo "
    "se descuenta del plazo."
)

# ══════════════════════════ del cliente ══════════════════════════
titulo("Lo que necesita el desarrollador")

vineta(
    "de cada prueba, con sus respuestas y resultados. Es la "
    "referencia contra la que se verifica que el sistema califique igual.",
    "Un protocolo real calificado a mano ",
)
vineta(
    "de dos horas para aprobar las claves de calificación. Puede "
    "designar a otro profesional si lo prefiere.",
    "Tres sesiones ",
)
vineta("de las licencias vigentes de las tres pruebas.", "Confirmación ")
vineta(
    "del sistema, entre $350 y $700 al mes, contratado a nombre del "
    "cliente. No va incluido en el precio.",
    "El alojamiento ",
)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(8)
p.add_run(
    "Mientras las claves no estén aprobadas por escrito, la plataforma muestra "
    "una advertencia y no se usa para emitir informes clínicos."
)

# ══════════════════════════ propiedad intelectual ══════════════════════════
doc.add_page_break()
titulo("Propiedad intelectual")

sub("Del cliente")
parrafo(
    "Los datos de los participantes, las respuestas, los resultados, los "
    "informes y las notas clínicas son propiedad exclusiva del cliente, desde "
    "el primer día y de forma permanente. El desarrollador no adquiere ningún "
    "derecho sobre esa información."
)

sub("Del desarrollador")
parrafo(
    "El código de la plataforma, su arquitectura, la implementación de los "
    "algoritmos de calificación, el diseño de las pantallas y el know-how "
    "empleado son propiedad del desarrollador, que puede reutilizar esos "
    "componentes en otros proyectos siempre que no incluyan información del "
    "cliente ni el contenido de las pruebas."
)

sub("La licencia del cliente")
parrafo(
    "El cliente recibe una licencia de uso perpetua, irrevocable, no "
    "exclusiva e intransferible sobre la plataforma, para el uso interno de su "
    "práctica profesional, sin límite de aplicaciones, participantes ni "
    "usuarios internos."
)
parrafo(
    "La licencia queda firme al completarse el pago total; antes de eso el "
    "permiso de uso es temporal, para prueba y validación. No comprende "
    "revender la plataforma, sublicenciarla, cederla ni prestar servicios de "
    "procesamiento a terceros con ella."
)

sub("Código fuente")
parrafo(
    "Con el pago total, el cliente recibe una copia del código fuente y puede "
    "modificarlo para su propio uso interno y para asegurar la continuidad de "
    "su operación. No se transfiere la propiedad ni se autoriza distribuirlo o "
    "comercializarlo."
)

sub("Si más adelante quiere más")
parrafo(
    "La propiedad del software, la exclusividad de uso o el derecho a "
    "comercializarlo se pueden acordar en un convenio separado, con una "
    "contraprestación distinta."
)

sub("Las pruebas")
parrafo(
    "PAPI, Hartman y MABE pertenecen a sus editores. Sus reactivos, claves, "
    "tablas y textos de interpretación no son propiedad de ninguna de las "
    "partes. El cliente manifiesta contar con las licencias para usarlas y "
    "digitalizarlas, y responde por cualquier reclamación derivada de su "
    "falta, dejando al desarrollador libre de responsabilidad."
)

# ══════════════════════════ acuerdos finales ══════════════════════════
titulo("Acuerdos finales")

sub("La herramienta")
parrafo(
    "Es un apoyo profesional, no un diagnóstico. Los textos salen marcados "
    "como borrador y ningún informe se emite sin la firma del cliente, a quien "
    "corresponden las decisiones clínicas y de selección. La responsabilidad "
    "del desarrollador no excede el monto pagado."
)

sub("Garantía")
parrafo(
    "Durante los treinta días siguientes a la última entrega, el desarrollador "
    "corrige sin costo los errores que se le reporten. Después de ese periodo, "
    "el mantenimiento y las mejoras se cotizan aparte."
)

sub("Reserva")
parrafo(
    "Lo que cada parte conozca de la otra queda reservado durante cinco años, "
    "en especial el contenido de las pruebas. El desarrollador puede mostrar "
    "capturas con datos ficticios para portafolio."
)

sub("Cierre")
parrafo(
    "El acuerdo termina al cumplirse el proyecto y cubrirse el pago, de común "
    "acuerdo, o por un incumplimiento que no se corrija dentro de los quince "
    "días siguientes al aviso. Si termina antes, se cubren las etapas "
    "concluidas y la parte proporcional de la que esté en curso, y la licencia "
    "no queda firme sobre lo no cubierto."
)
parrafo(
    "Las diferencias se conversan primero. De no resolverse, aplican las leyes "
    "de [estado] y los tribunales de [ciudad]. Los cambios a este acuerdo van "
    "por escrito."
)

# ══════════════════════════ firmas ══════════════════════════
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(24)
p.add_run("Ambas partes leyeron y aceptaron. Se firma por duplicado.")

doc.add_paragraph()
doc.add_paragraph()

t = doc.add_table(rows=3, cols=2)
t.columns[0].width = Cm(7.9)
t.columns[1].width = Cm(7.9)
firmas = [
    ("______________________________", "______________________________"),
    ("[nombre]", "[nombre]"),
    ("El desarrollador", "El cliente"),
]
for i, (fila, (a, b)) in enumerate(zip(t.rows, firmas)):
    for celda, txt in zip(fila.cells, (a, b)):
        celda.width = Cm(7.9)
        pp = celda.paragraphs[0]
        pp.paragraph_format.space_after = Pt(2)
        r = pp.add_run(txt)
        if i == 2:
            r.font.size = Pt(9)
            r.font.color.rgb = SUAVE

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(16)
r = p.add_run("Los datos de facturación se intercambian por separado.")
r.font.size = Pt(9)
r.font.color.rgb = SUAVE

# ══════════════════════════ anexo A ══════════════════════════
doc.add_page_break()
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(12)
r = p.add_run("Anexo A · Alcance")
r.font.size = Pt(16)
r.font.name = "Cambria"

sub("Aplicación")
for x in [
    "Instrucciones de cada prueba, tomadas de sus manuales.",
    "Aviso de privacidad y aceptación de términos antes de iniciar, con registro de fecha y versión.",
    "Validación al capturar: no permite guardar un protocolo incompleto, repetido o fuera de rango.",
    "Guardado automático, con opción de interrumpir y retomar.",
    "Registro de participantes, sesiones y perfiles de puesto.",
]:
    vineta(x)

sub("Calificación")
for x in [
    "Puntajes automáticos de las tres pruebas a partir de las respuestas.",
    "Puertas de validez: sumas de control, paridad de disimilitudes y bloqueo de la interpretación cuando el instrumento lo determina.",
    "Desglose reactivo por reactivo.",
]:
    vineta(x)

sub("Interpretación e informes")
for x in [
    "Hoja de perfil del PAPI, axiogramas del Hartman y gráficas del MABE.",
    "Textos compuestos a partir de los manuales.",
    "Panel de revisión: listado con estado, detalle de resultados, historial por participante, notas, aprobación y firma.",
    "Informe en PDF con fecha, responsable y versión de claves. Una vez firmado, queda inalterable.",
]:
    vineta(x)

sub("Seguridad y arranque")
for x in [
    "Tres perfiles: aplicador, psicólogo y administrador. El administrador no ve contenido clínico.",
    "Acceso con contraseña, tráfico cifrado, base de datos y respaldos cifrados, bitácora de operaciones.",
    "Respaldos automáticos, con una restauración probada antes de entregar.",
    "Instalación, puesta en marcha, capacitación y documentación.",
]:
    vineta(x)

sub("No incluido")
parrafo(
    "Alojamiento y dominio · aplicación en pantalla para que el evaluado "
    "conteste directamente · app móvil · integración con otros sistemas · "
    "pruebas distintas a las tres señaladas · revisión legal del aviso de "
    "privacidad · mantenimiento posterior a la garantía."
)

# ══════════════════════════ anexo B ══════════════════════════
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(22)
p.paragraph_format.space_after = Pt(12)
r = p.add_run("Anexo B · Etapas")
r.font.size = Pt(16)
r.font.name = "Cambria"

etapas = [
    (
        "Primera etapa",
        "$5,000 + IVA",
        "Claves de calificación validadas contra los protocolos del cliente y "
        "aprobadas por él. La plataforma captura y guarda las tres pruebas, con "
        "usuarios y perfiles funcionando.",
    ),
    (
        "Segunda etapa",
        "$5,000 + IVA",
        "Las tres pruebas calculan sus puntajes a partir de las respuestas, con "
        "las puertas de validez operando, y el panel permite revisarlos y "
        "aprobarlos.",
    ),
    (
        "Tercera etapa",
        "cierre",
        "Gráficas e informes en PDF con trazabilidad, sistema instalado en "
        "producción, capacitación impartida y documentación entregada.",
    ),
]
for nombre, monto, texto in etapas:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(nombre)
    r.font.bold = True
    r2 = p.add_run("     " + monto)
    r2.font.size = Pt(9)
    r2.font.color.rgb = SUAVE
    parrafo(texto, sangria=0.7)

sub("Insumos del cliente")
t = doc.add_table(rows=4, cols=2)
t.columns[0].width = Cm(4.6)
t.columns[1].width = Cm(11.2)
insumos = [
    ("Antes de arrancar", "Un protocolo calificado a mano de cada prueba y la confirmación de las licencias vigentes."),
    ("Antes del mes 2", "Tablas normativas que tenga en su poder."),
    ("Antes del mes 3", "Textos de interpretación de los manuales, en la plantilla proporcionada."),
    ("Meses 1, 2 y 3", "Disponibilidad para las sesiones de validación."),
]
for fila, (k, v) in zip(t.rows, insumos):
    fila.cells[0].width = Cm(4.6)
    fila.cells[1].width = Cm(11.2)
    pk = fila.cells[0].paragraphs[0]
    pk.paragraph_format.space_after = Pt(5)
    rk = pk.add_run(k)
    rk.font.size = Pt(9)
    rk.font.color.rgb = SUAVE
    pv = fila.cells[1].paragraphs[0]
    pv.paragraph_format.space_after = Pt(5)
    pv.add_run(v)

doc.save("Acuerdo de colaboración.docx")
print("Acuerdo de colaboración.docx")
